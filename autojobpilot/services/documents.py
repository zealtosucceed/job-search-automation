"""Document generation (PRD Features 7, 8, 10).

Customized CV, cover letter, and LinkedIn outreach message. DOCX + Markdown only
(no PDF). The CV uses only truthful source material — the prompt forbids
fabrication; we just render whatever structured CV the LLM returns.
"""

from __future__ import annotations

import json
from pathlib import Path

from docx import Document

from .. import prompts
from ..models import FitAnalysis, Job
from ..utils import get_logger
from .llm.base import LLMAdapter
from .scorer import job_description_text

log = get_logger("document_generation")


# ── CV ──────────────────────────────────────────────────────────────────────
def generate_cv(job: Job, resume_text: str, analysis: FitAnalysis,
                out_dir: Path, llm: LLMAdapter) -> dict[str, str]:
    prompt = prompts.render(
        "cv_customizer",
        resume_text=resume_text,
        job_description=job_description_text(job),
        fit_analysis=json.dumps(analysis.to_dict(), indent=2),
    )
    cv = llm.complete_json(prompt)
    if not isinstance(cv, dict):
        raise ValueError("CV customizer did not return a JSON object")

    docx_path = out_dir / "custom_cv.docx"
    md_path = out_dir / "custom_cv.md"
    _write_cv_docx(cv, docx_path)
    md_path.write_text(_cv_markdown(cv), encoding="utf-8")
    log.info("CV generated: %s", docx_path)
    return {"docx": str(docx_path), "md": str(md_path)}


def _write_cv_docx(cv: dict, path: Path) -> None:
    doc = Document()
    name = cv.get("name", "")
    if name:
        doc.add_heading(name, level=0)
    if cv.get("contact"):
        doc.add_paragraph(cv["contact"])
    if cv.get("professional_summary"):
        doc.add_heading("Professional Summary", level=1)
        doc.add_paragraph(cv["professional_summary"])
    if cv.get("skills"):
        doc.add_heading("Skills", level=1)
        doc.add_paragraph(", ".join(cv["skills"]))
    if cv.get("experience"):
        doc.add_heading("Experience", level=1)
        for exp in cv["experience"]:
            head = " — ".join(x for x in (exp.get("role"), exp.get("company")) if x)
            p = doc.add_paragraph()
            p.add_run(head).bold = True
            if exp.get("duration"):
                p.add_run(f"  ({exp['duration']})")
            for bullet in exp.get("bullets", []):
                doc.add_paragraph(bullet, style="List Bullet")
    if cv.get("education"):
        doc.add_heading("Education", level=1)
        for ed in cv["education"]:
            doc.add_paragraph(ed if isinstance(ed, str) else json.dumps(ed))
    if cv.get("certifications"):
        doc.add_heading("Certifications", level=1)
        for c in cv["certifications"]:
            doc.add_paragraph(c if isinstance(c, str) else json.dumps(c), style="List Bullet")
    for sec in cv.get("additional_sections", []) or []:
        if isinstance(sec, dict):
            doc.add_heading(sec.get("title", "Additional"), level=1)
            for item in sec.get("items", []):
                doc.add_paragraph(str(item), style="List Bullet")
    doc.save(str(path))


def _cv_markdown(cv: dict) -> str:
    lines: list[str] = []
    if cv.get("name"):
        lines.append(f"# {cv['name']}")
    if cv.get("contact"):
        lines.append(cv["contact"])
    if cv.get("professional_summary"):
        lines += ["", "## Professional Summary", cv["professional_summary"]]
    if cv.get("skills"):
        lines += ["", "## Skills", ", ".join(cv["skills"])]
    if cv.get("experience"):
        lines += ["", "## Experience"]
        for exp in cv["experience"]:
            head = " — ".join(x for x in (exp.get("role"), exp.get("company")) if x)
            dur = f" ({exp['duration']})" if exp.get("duration") else ""
            lines.append(f"### {head}{dur}")
            lines += [f"- {b}" for b in exp.get("bullets", [])]
    if cv.get("education"):
        lines += ["", "## Education"]
        lines += [f"- {ed if isinstance(ed, str) else json.dumps(ed)}" for ed in cv["education"]]
    if cv.get("certifications"):
        lines += ["", "## Certifications"]
        lines += [f"- {c if isinstance(c, str) else json.dumps(c)}" for c in cv["certifications"]]
    return "\n".join(lines) + "\n"


# ── Cover letter ────────────────────────────────────────────────────────────
def generate_cover_letter(job: Job, resume_text: str, analysis: FitAnalysis,
                          out_dir: Path, llm: LLMAdapter) -> dict[str, str]:
    prompt = prompts.render(
        "cover_letter",
        resume_text=resume_text,
        job_description=job_description_text(job),
        fit_analysis=json.dumps(analysis.to_dict(), indent=2),
    )
    data = llm.complete_json(prompt)
    text = data.get("cover_letter", "") if isinstance(data, dict) else str(data)

    docx_path = out_dir / "cover_letter.docx"
    md_path = out_dir / "cover_letter.md"
    doc = Document()
    for para in text.split("\n\n"):
        doc.add_paragraph(para.strip())
    doc.save(str(docx_path))
    md_path.write_text(text + "\n", encoding="utf-8")
    log.info("Cover letter generated: %s", docx_path)
    return {"docx": str(docx_path), "md": str(md_path)}


# ── LinkedIn outreach ───────────────────────────────────────────────────────
def generate_outreach(job: Job, resume_text: str, out_dir: Path,
                      llm: LLMAdapter) -> dict[str, str] | None:
    """Only generated when a contact is present (PRD Feature 10)."""
    contact = job.recruiter_name
    contact_title = ""
    if not contact and job.hiring_team_contacts:
        first = job.hiring_team_contacts[0]
        contact = first.get("contact_name")
        contact_title = first.get("contact_title", "")
    if not contact:
        return None

    prompt = prompts.render(
        "linkedin_message",
        resume_text=resume_text,
        job_description=job_description_text(job),
        company=job.company,
        job_title=job.job_title,
        contact_name=contact,
        contact_title=contact_title,
    )
    data = llm.complete_json(prompt)
    msg = data.get("linkedin_message", "") if isinstance(data, dict) else str(data)
    txt_path = out_dir / "linkedin_message.txt"
    txt_path.write_text(msg + "\n", encoding="utf-8")
    log.info("Outreach message generated: %s", txt_path)
    return {"txt": str(txt_path)}
